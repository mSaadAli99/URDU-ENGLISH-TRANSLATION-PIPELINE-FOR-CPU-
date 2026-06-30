# ============================================================
# pipeline/export.py — Stage 6: Final Export
# Generates:
#   1. Structured JSON dataset
#   2. Professional DOCX report
# ============================================================

import os
import json
from datetime import datetime
from pipeline.utils import save_json, ensure_dirs, print_banner, now_str

import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import config


def export(stage5_result: dict) -> dict:
    """
    Stage 6: Export final de-identified dataset as JSON + DOCX.

    Args:
        stage5_result: Output dict from Stage 5 (deidentify.py)

    Returns:
        dict with paths to the two generated output files
    """
    print_banner(6, "FINAL EXPORT: JSON + DOCX")

    interview_id = stage5_result["interview_id"]
    ensure_dirs(config.STAGE6_DIR)

    print(f"  Interview ID : {interview_id}")

    # ── 1. Build final JSON ───────────────────────────────────
    print("\n  Building final JSON dataset...")

    # Collect all flagged segments from both verification stages
    transcript_flags   = stage5_result["transcript_verification_report"].get("flagged_details", [])
    translation_flags  = stage5_result["translation_verification_report"].get("flagged_details", [])
    all_flags = transcript_flags + translation_flags

    final_json = {
        # ── Metadata ─────────────────────────────────
        "interview_id"                  : interview_id,
        "audio_filename"                : stage5_result["audio_filename"],
        "duration_minutes"              : stage5_result["duration_minutes"],
        "processing_date"               : now_str(),
        "pipeline_version"              : "1.0.0",

        # ── Quality scores ────────────────────────────
        "transcript_quality_score"      : stage5_result["transcript_quality_score"],
        "translation_quality_score"     : stage5_result["translation_quality_score"],

        # ── Text at each stage ────────────────────────
        "urdu_transcript"               : stage5_result["urdu_full_text"],
        "verified_urdu_transcript"      : stage5_result.get("verified_urdu_full", ""),
        "english_translation"           : stage5_result["english_full_text"],
        "verified_english_translation"  : stage5_result["verified_english_full"],
        "deidentified_english"          : stage5_result["deidentified_english_full"],

        # ── De-identification summary ──────────────────
        "entities_removed"              : stage5_result["unique_entities_removed"],
        "entities_removed_count"        : stage5_result["entities_removed_count"],
        "entities_removed_by_type"      : stage5_result["entities_removed_by_type"],

        # ── Flagged segments ──────────────────────────
        "flagged_segments"              : all_flags,

        # ── Verification reports ──────────────────────
        "transcript_verification_report" : stage5_result["transcript_verification_report"],
        "translation_verification_report": stage5_result["translation_verification_report"],

        # ── Full segment-level data ───────────────────
        "segments"                      : stage5_result["deidentified_segments"],
    }

    json_path = os.path.join(config.STAGE6_DIR, f"{interview_id}_final_dataset.json")
    save_json(final_json, json_path)

    # ── 2. Build final DOCX ───────────────────────────────────
    print("\n  Building DOCX report...")
    docx_path = _build_docx(final_json, interview_id)

    # ── Summary ───────────────────────────────────────────────
    print(f"\n  ── Final Outputs ─────────────────────────────")
    print(f"  JSON → {json_path}")
    print(f"  DOCX → {docx_path}")

    result = {
        "stage"        : 6,
        "stage_name"   : "Final Export",
        "interview_id" : interview_id,
        "processed_at" : now_str(),
        "json_path"    : json_path,
        "docx_path"    : docx_path,
        "final_dataset": final_json,
    }

    print(f"\n  ✔ Stage 6 complete. Pipeline finished!")
    return result


def _build_docx(data: dict, interview_id: str) -> str:
    """Build a professional DOCX report from the final dataset."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("  ⚠ python-docx not installed. Skipping DOCX export.")
        print("    Run: pip install python-docx")
        return ""

    doc = Document()

    # ── Set margins ───────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Helper: add styled heading ────────────────────────────
    def add_heading(text, level=1, color=(31, 73, 125)):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
        return h

    # ── Helper: add key-value row in a table ──────────────────
    def add_kv_row(table, key, value):
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True

    # ════════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════════
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(config.DOCX_TITLE)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(31, 73, 125)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Interview ID: {interview_id}").font.size = Pt(13)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Processed: {data['processing_date'][:10]}").font.size = Pt(11)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # SECTION 1: METADATA
    # ════════════════════════════════════════════════════════
    add_heading("1. Interview Metadata", level=1)

    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.style = "Light Shading Accent 1"
    meta_table.rows[0].cells[0].text = "Field"
    meta_table.rows[0].cells[1].text = "Value"
    for cell in meta_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    meta_fields = [
        ("Interview ID",             data["interview_id"]),
        ("Audio File",               data["audio_filename"]),
        ("Duration (minutes)",       data["duration_minutes"]),
        ("Processing Date",          data["processing_date"][:10]),
        ("Transcript Quality Score", f"{data['transcript_quality_score']}/100"),
        ("Translation Quality Score",f"{data['translation_quality_score']}/100"),
        ("Entities Removed",         data["entities_removed_count"]),
        ("Flagged Segments",         len(data["flagged_segments"])),
    ]
    for k, v in meta_fields:
        add_kv_row(meta_table, k, v)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════
    # SECTION 2: VERIFIED URDU TRANSCRIPT
    # ════════════════════════════════════════════════════════
    add_heading("2. Verified Urdu Transcript", level=1)
    doc.add_paragraph(
        "Original Urdu transcription with low-confidence segments marked as [LOW CONFIDENCE]."
    ).italic = True

    urdu_para = doc.add_paragraph()
    urdu_para.add_run(data.get("verified_urdu_transcript", data["urdu_transcript"]))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # SECTION 3: ENGLISH TRANSLATION
    # ════════════════════════════════════════════════════════
    add_heading("3. English Translation", level=1)
    doc.add_paragraph(
        "Verified English translation. Segments that failed language detection are marked as [TRANSLATION FAILED]."
    ).italic = True

    doc.add_paragraph(data.get("verified_english_translation", data["english_translation"]))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # SECTION 4: FINAL DE-IDENTIFIED ENGLISH TEXT
    # ════════════════════════════════════════════════════════
    add_heading("4. Final De-identified English Text", level=1)
    doc.add_paragraph(
        "All personally identifiable information (PII) has been removed and replaced with placeholder tags."
    ).italic = True

    doc.add_paragraph(data["deidentified_english"])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # SECTION 5: ENTITIES REMOVED
    # ════════════════════════════════════════════════════════
    add_heading("5. Entities Removed (De-identification Summary)", level=1)

    # By type counts
    add_heading("5.1 Summary by Type", level=2)
    if data["entities_removed_by_type"]:
        ent_table = doc.add_table(rows=1, cols=2)
        ent_table.style = "Light Shading Accent 1"
        ent_table.rows[0].cells[0].text = "Entity Type"
        ent_table.rows[0].cells[1].text = "Count"
        for cell in ent_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
        for etype, count in sorted(data["entities_removed_by_type"].items(), key=lambda x: -x[1]):
            add_kv_row(ent_table, etype, count)
    else:
        doc.add_paragraph("No PII entities detected in this interview.")

    doc.add_paragraph()

    # Full entity list
    add_heading("5.2 Full Entity List", level=2)
    if data["entities_removed"]:
        full_ent_table = doc.add_table(rows=1, cols=3)
        full_ent_table.style = "Light Shading Accent 1"
        headers = ["Original Text", "Type", "Replaced With"]
        for i, h in enumerate(headers):
            full_ent_table.rows[0].cells[i].text = h
            full_ent_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for ent in data["entities_removed"]:
            row = full_ent_table.add_row()
            row.cells[0].text = ent.get("original", "")
            row.cells[1].text = ent.get("type", "")
            row.cells[2].text = ent.get("replaced_with", "")
    else:
        doc.add_paragraph("No PII entities found.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # SECTION 6: VERIFICATION REPORTS
    # ════════════════════════════════════════════════════════
    add_heading("6. Verification Reports", level=1)

    # Transcript verification
    add_heading("6.1 Transcript Verification Report", level=2)
    tr = data["transcript_verification_report"]
    tr_table = doc.add_table(rows=1, cols=2)
    tr_table.style = "Light Shading Accent 1"
    tr_table.rows[0].cells[0].text = "Metric"
    tr_table.rows[0].cells[1].text = "Value"
    tr_fields = [
        ("Total Segments",      tr.get("total_segments", "")),
        ("Good Segments",       tr.get("good_segments", "")),
        ("Flagged Segments",    tr.get("flagged_segments", "")),
        ("Avg Confidence",      tr.get("avg_confidence", "")),
        ("Quality Score",       f"{tr.get('quality_score', '')}/100"),
        ("Quality Label",       tr.get("quality_label", "")),
        ("Duration Covered",    f"{tr.get('duration_covered_min', '')} min"),
    ]
    for k, v in tr_fields:
        add_kv_row(tr_table, k, v)

    doc.add_paragraph()

    # Translation verification
    add_heading("6.2 Translation Verification Report", level=2)
    tvr = data["translation_verification_report"]
    tv_table = doc.add_table(rows=1, cols=2)
    tv_table.style = "Light Shading Accent 1"
    tv_table.rows[0].cells[0].text = "Metric"
    tv_table.rows[0].cells[1].text = "Value"
    tv_fields = [
        ("Total Segments",   tvr.get("total_segments", "")),
        ("Good Segments",    tvr.get("good_segments", "")),
        ("Flagged Segments", tvr.get("flagged_segments", "")),
        ("Quality Score",    f"{tvr.get('overall_quality_score', '')}/100"),
        ("Quality Label",    tvr.get("quality_label", "")),
    ]
    for k, v in tv_fields:
        add_kv_row(tv_table, k, v)

    # ── Save DOCX ─────────────────────────────────────────────
    docx_path = os.path.join(config.STAGE6_DIR, f"{interview_id}_final_dataset.docx")
    doc.save(docx_path)
    print(f"  ✔ DOCX saved → {docx_path}")
    return docx_path
